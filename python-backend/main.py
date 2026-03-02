"""
The Forge - Python Backend Service
Handles document export (DOCX, PDF) and text processing utilities.
"""

from fastapi import FastAPI, HTTPException, Body
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import io
import json

# Import document generation libraries
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY

app = FastAPI(
    title="The Forge Python Backend",
    description="Export and text processing service for The Forge writing platform",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "https://theauditorsforge.com"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class SceneData(BaseModel):
    chapterTitle: str
    sceneTitle: str
    content: str
    wordCount: Optional[int] = 0


class ExportRequest(BaseModel):
    title: str
    author: Optional[str] = "The Author"
    scenes: List[SceneData]
    format: str  # "docx" or "pdf"


class TextAnalysisRequest(BaseModel):
    text: str


class TextAnalysisResponse(BaseModel):
    wordCount: int
    sentenceCount: int
    paragraphCount: int
    avgWordsPerSentence: float
    shortSentenceRatio: float
    longSentenceRatio: float
    passiveVoiceEstimate: float
    dialogueRatio: float
    uniqueWordRatio: float
    readingTimeMinutes: float


@app.get("/health")
async def health_check():
    return {"status": "ok", "service": "forge-python-backend", "version": "1.0.0"}


@app.post("/export/docx")
async def export_docx(request: ExportRequest = Body(...)):
    """Export manuscript as .docx file"""
    try:
        doc = DocxDocument()

        # Document settings
        for section in doc.sections:
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(request.title.upper())
        title_run.bold = True
        title_run.font.size = Pt(16)

        if request.author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"by {request.author}")
            author_run.font.size = Pt(12)
            author_run.italic = True

        doc.add_page_break()

        current_chapter = ""
        for scene in request.scenes:
            # Chapter header
            if scene.chapterTitle != current_chapter:
                current_chapter = scene.chapterTitle
                ch_para = doc.add_paragraph()
                ch_run = ch_para.add_run(scene.chapterTitle)
                ch_run.bold = True
                ch_run.font.size = Pt(14)
                ch_para.space_after = Pt(12)

            # Scene title
            sc_para = doc.add_paragraph()
            sc_run = sc_para.add_run(scene.sceneTitle)
            sc_run.bold = True
            sc_run.font.size = Pt(12)
            sc_para.space_after = Pt(6)

            # Scene content
            if scene.content:
                paragraphs = [p.strip() for p in scene.content.split("\n\n") if p.strip()]
                for para_text in paragraphs:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = p.add_run(para_text)
                    run.font.size = Pt(12)
                    p.space_after = Pt(6)
                    p.paragraph_format.first_line_indent = Pt(18)

            # Scene break
            break_para = doc.add_paragraph()
            break_run = break_para.add_run("* * *")
            break_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break_run.italic = True
            break_para.space_before = Pt(12)
            break_para.space_after = Pt(12)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"{request.title.replace(' ', '_')}.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")


@app.post("/export/pdf")
async def export_pdf(request: ExportRequest = Body(...)):
    """Export manuscript as .pdf file"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=1.25 * inch,
            leftMargin=1.25 * inch,
            topMargin=1 * inch,
            bottomMargin=1 * inch
        )

        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        title_style = ParagraphStyle(
            "ForgeTitle",
            parent=styles["Title"],
            fontSize=18,
            fontName="Helvetica-Bold",
            spaceAfter=12,
            alignment=TA_JUSTIFY,
        )
        chapter_style = ParagraphStyle(
            "ForgeChapter",
            parent=styles["Heading1"],
            fontSize=14,
            fontName="Helvetica-Bold",
            spaceAfter=8,
            spaceBefore=20,
        )
        scene_style = ParagraphStyle(
            "ForgeScene",
            parent=styles["Heading2"],
            fontSize=12,
            fontName="Helvetica-BoldOblique",
            spaceAfter=6,
            spaceBefore=12,
        )
        body_style = ParagraphStyle(
            "ForgeBody",
            parent=styles["Normal"],
            fontSize=12,
            fontName="Helvetica",
            leading=18,
            spaceAfter=6,
            firstLineIndent=18,
            alignment=TA_JUSTIFY,
        )
        break_style = ParagraphStyle(
            "ForgeBreak",
            parent=styles["Normal"],
            fontSize=12,
            fontName="Helvetica-Oblique",
            alignment=1,  # CENTER
            spaceAfter=12,
            spaceBefore=12,
        )

        # Title page
        story.append(Paragraph(request.title.upper(), title_style))
        if request.author:
            story.append(Paragraph(f"by {request.author}", styles["Normal"]))
        story.append(Spacer(1, 0.5 * inch))

        current_chapter = ""
        for scene in request.scenes:
            if scene.chapterTitle != current_chapter:
                current_chapter = scene.chapterTitle
                story.append(Paragraph(scene.chapterTitle, chapter_style))

            story.append(Paragraph(scene.sceneTitle, scene_style))

            if scene.content:
                paragraphs = [p.strip() for p in scene.content.split("\n\n") if p.strip()]
                for para_text in paragraphs:
                    story.append(Paragraph(para_text, body_style))

            story.append(Paragraph("* * *", break_style))

        doc.build(story)
        buffer.seek(0)

        filename = f"{request.title.replace(' ', '_')}.pdf"

        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")


@app.post("/analyze/text", response_model=TextAnalysisResponse)
async def analyze_text(request: TextAnalysisRequest):
    """Perform statistical text analysis for Temperature Check"""
    import re

    text = request.text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="Text is required")

    words = text.split()
    word_count = len(words)

    sentences = re.split(r'[.!?]+', text)
    sentences = [s.strip() for s in sentences if s.strip()]
    sentence_count = len(sentences)

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    paragraph_count = len(paragraphs)

    avg_words = word_count / max(sentence_count, 1)

    sentence_lengths = [len(s.split()) for s in sentences]
    short_sentences = sum(1 for l in sentence_lengths if l < 8)
    long_sentences = sum(1 for l in sentence_lengths if l > 30)

    short_ratio = short_sentences / max(sentence_count, 1)
    long_ratio = long_sentences / max(sentence_count, 1)

    # Passive voice detection (regex-based approximation)
    passive_pattern = re.compile(r'\b(is|are|was|were|be|been|being)\s+\w+ed\b', re.IGNORECASE)
    passive_count = len(passive_pattern.findall(text))
    passive_ratio = passive_count / max(sentence_count, 1)

    # Dialogue ratio
    dialogue_matches = re.findall(r'["""].*?["""]', text, re.DOTALL)
    dialogue_chars = sum(len(m) for m in dialogue_matches)
    dialogue_ratio = dialogue_chars / max(len(text), 1)

    # Lexical diversity
    unique_words = len(set(w.lower().strip('.,!?"\'') for w in words if w.strip('.,!?"\'').isalpha()))
    unique_ratio = unique_words / max(word_count, 1)

    # Reading time (average 238 wpm for adults)
    reading_time = word_count / 238.0

    return TextAnalysisResponse(
        wordCount=word_count,
        sentenceCount=sentence_count,
        paragraphCount=paragraph_count,
        avgWordsPerSentence=round(avg_words, 2),
        shortSentenceRatio=round(short_ratio, 3),
        longSentenceRatio=round(long_ratio, 3),
        passiveVoiceEstimate=round(passive_ratio, 3),
        dialogueRatio=round(dialogue_ratio, 3),
        uniqueWordRatio=round(unique_ratio, 3),
        readingTimeMinutes=round(reading_time, 1),
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001, reload=True)
